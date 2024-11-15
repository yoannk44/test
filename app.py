import streamlit as st
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io
import os

# Configuration du port pour Heroku
port = int(os.environ.get("PORT", 8501))

def create_word_report(image_path, title="Rapport d'analyse", author="Auteur"):
    # Création du document
    doc = Document()
    
    # Style du titre principal
    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run(title)
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Ajout de la date
    date_paragraph = doc.add_paragraph()
    date_run = date_paragraph.add_run(f"Date: {datetime.now().strftime('%d/%m/%Y')}")
    date_run.font.size = Pt(10)
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Ajout de l'auteur
    author_paragraph = doc.add_paragraph()
    author_run = author_paragraph.add_run(f"Auteur: {author}")
    author_run.font.size = Pt(10)
    author_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Ajout d'une ligne de séparation
    doc.add_paragraph("_" * 50)
    
    # Section Introduction
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("Ce rapport présente l'analyse des données visuelles collectées.")
    
    # Section Image
    doc.add_heading("Analyse Visuelle", level=1)
    doc.add_paragraph("Ci-dessous, vous trouverez l'image analysée:")
    
    # Ajout de l'image
    doc.add_picture(image_path, width=Inches(6))
    
    # Légende de l'image
    caption = doc.add_paragraph("Figure 1: Visualisation des données")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Section Analyse
    doc.add_heading("Analyse Détaillée", level=1)
    doc.add_paragraph("L'analyse de l'image ci-dessus révèle les points suivants:")
    
    # Ajout d'une liste à puces
    points = [
        "Point d'analyse 1",
        "Point d'analyse 2",
        "Point d'analyse 3"
    ]
    for point in points:
        paragraph = doc.add_paragraph()
        paragraph.add_run("• ").font.bold = True
        paragraph.add_run(point)
    
    # Section Conclusion
    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph("En conclusion, cette analyse nous permet de...")
    
    # Sauvegarder le document dans un buffer
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    
    return doc_buffer

def main():
    st.title("Générateur de Rapport Word avec Images")
    
    # Upload de l'image
    uploaded_file = st.file_uploader("Choisissez une image", type=["jpg", "jpeg", "png"])
    
    # Champs de personnalisation
    title = st.text_input("Titre du rapport", "Rapport d'analyse")
    author = st.text_input("Auteur du rapport", "Auteur")
    
    if uploaded_file is not None:
        # Afficher l'aperçu de l'image
        st.image(uploaded_file, caption="Aperçu de l'image", use_column_width=True)
        
        # Bouton pour générer le rapport
        if st.button("Générer le rapport Word"):
            doc_buffer = create_word_report(uploaded_file, title, author)
            
            # Téléchargement du document
            st.download_button(
                label="Télécharger le rapport",
                data=doc_buffer,
                file_name=f"rapport_{datetime.now().strftime('%Y%m%d')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

if __name__ == "__main__":
    main()
